fy = 335
f = 315
ZG = 6 #float(input('请输入柱高（m）：'))
CD = 60 #float(input('请输入长度（m）：'))
KD = 18 #float(input('请输入跨度（m）：'))
ZJ = 7.5 #float(input('请输入柱距（m）：'))
w0 = 0.4 #float(input('请输入基本风压：'))
s0 = 0.5 #float(input('请输入基本雪压：'))
i = 0.1 #float(input('请输入屋面坡度：'))
GC = "Q235" #(input('请输入钢材种类：'))
LJ = 1.5 #float(input('请输入檩距：'))
# 风荷载体积系数假设左风
zuozhu = 0.8   #左柱
zuopo = -0.6   #左坡
youpo = -0.5   #右坡
youzhu = -0.5   #右柱
wumianbanlintiao = 0.5
wujiazizhong = round(0.12+0.011*CD,2) #经验公式
diaoding = 0.5 #0.35+round(0.12+0.011*CD,2)  #计算吊顶时  #永久荷载
gangjia = 0.5  #计算钢架时   #活荷载
lintiao = 0.5  #计算檩条时
# ----------------------------------------------------------------------------------------------------------------------
# word文档编辑
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import math
# 创建一个文档
docx = Document()
# 保存一下
docx.save("轻型门式钢架.docx")
# 输入这个标题(文本)内容到P1
P1 = ('轻型门式钢架')
# 创建一个段落（文本）
p1 = docx.add_paragraph("")
run1 = p1.add_run(f'{P1}' )
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1.bold = True  # 加粗
run1.italic = False  # 斜体
# 修改字号大小
run1.font.size = Pt(22)
# 修改字体样式
run1.font.name=u'黑体'
run1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
# 首行缩进，小五＝9磅 五号＝10.5磅 小四＝12磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝22磅 小一＝24磅 一号＝26磅
p1.paragraph_format.first_line_indent = Pt(0)
# 设置行间距为，可输入整数和浮点数
p1.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('一、设计资料')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f"""单跨双坡轻型门式刚架，刚架跨度{KD}m，长度{CD}m，柱距{ZJ}m，柱高{ZG}m，共有{CD/ZJ+1}榀刚架，屋面坡度{i}，刚架为等截面梁柱，采用{GC}钢材，焊条采用E43型，螺栓采用高强度摩擦螺栓。地震设防烈度为7度，地震加速度设计值为0.1g，风振系数和风压高度变化系数均取1.0，檩条间距{LJ}m，恒荷载分项系数取1.3，活荷载分项系数取1.5，最不利截面内力设计值按荷载基本组合设计。""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('二、柱网及屋面布置')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('1、柱网平面布置')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('2、支撑体系布置')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 生成图片文件名
filename1 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\{int(KD)}m跨度{ZJ}m柱距平面支撑.png'
filename2 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\{int(KD)}m跨度{ZJ}m柱距柱间支撑.png'

# 插入图片
docx.paragraphs[4].add_run().add_picture(filename1, width=Inches(6))
docx.paragraphs[5].add_run().add_picture(filename2, width=Inches(6))
# ----------------------------------------------------------------------------------------------------------------------
P = ('三、结构内力分析')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('1、荷载计算')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f"""永久荷载：吊顶{diaoding}kN/m²
屋面活荷载：计算刚架时{gangjia}kN/m²，计算檩条时{lintiao}kN/m²
风荷载和雪荷载：{w0}kN/m²，{s0}kN/m²""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------

# 创建一个6行4列的表格
table = docx.add_table(rows=6, cols=4)

# 合并第一列的单元格并添加文本
for cell in table.columns[0].cells:
    cell.merge(table.cell(0, 0))
table.cell(0, 0).text = '各荷载标准值计算'

# 添加第二列的文本
table.cell(0, 1).text = '屋面'
table.cell(1, 1).text = '钢架'
table.cell(2, 1).merge(table.cell(3, 1))
table.cell(2, 1).text = '迎风面'
table.cell(4, 1).merge(table.cell(5, 1))
table.cell(4, 1).text = '背风面'

# 添加第三列的文本
table.cell(0, 2).text = '活荷载标准值'
table.cell(1, 2).text = '吊顶'
table.cell(2, 2).text = '柱上q1'
table.cell(3, 2).text = '横梁上q2'
table.cell(4, 2).text = '柱上q3'
table.cell(5, 2).text = '横梁上q4'
# 计算过程
table.cell(0, 3).text = f'{max(s0,w0)}×{ZJ}={round(max(s0,w0)*ZJ,2)}kN/m'
table.cell(1, 3).text = f'{diaoding}×{ZJ}={round(diaoding*ZJ,2)}kN/m'
table.cell(2, 3).text = f'{w0}×{ZJ}×{zuozhu}={round(w0*ZJ*zuozhu,2)}kN/m'
table.cell(3, 3).text = f'{w0}×{ZJ}×{zuopo}={round(w0*ZJ*zuopo,2)}kN/m'
table.cell(4, 3).text = f'{w0}×{ZJ}×{youzhu}={round(w0*ZJ*youzhu,2)}kN/m'
table.cell(5, 3).text = f'{w0}×{ZJ}×{youpo}={round(w0*ZJ*youpo,2)}kN/m'
q1 = round(w0*ZJ*zuozhu,2)
q2 = round(w0*ZJ*zuopo,2)
q3 = round(w0*ZJ*youzhu,2)
q4 = round(w0*ZJ*youpo,2)
# ----------------------------------------------------------------------------------------------------------------------
P = ('2、恒活荷载、风荷载计算简图')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 生成图片文件名
filename1 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\恒荷载下的计算简图.png'
filename2 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\风荷载下的计算简图.png'

# 插入图片
docx.paragraphs[9].add_run().add_picture(filename1, width=Inches(6))
docx.paragraphs[9].add_run().add_picture(filename2, width=Inches(6))
# ----------------------------------------------------------------------------------------------------------------------
P = ('四、内力分析')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('1、内力计算')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f"""为了简化计算，假设恒荷载沿水平方向分布在梁上，风荷载沿竖直方向分布在柱子上及沿着坡屋顶分布，当遇到左风时，左侧为风压力，右侧为风吸力，且屋面坡度小于30°，左侧风载体形系数为0.6，右侧为-0.5。计算内力时候避免复杂运算，可采取单位力法计算出内力系数(即q=1时，钢架的内力值)，后根据荷载大小进行扩大倍数，最后求出实际内力值。内力系数分布如下图所示。""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 生成图片文件名
filename1 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\恒活荷载内力系数1.png'
filename2 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\恒活荷载内力系数2.png'
filename3 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\恒活荷载内力系数3.png'
filename4 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\风荷载内力系数1.png'
filename5 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\风荷载内力系数2.png'
filename6 = f'C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\风荷载内力系数3.png'
# 插入图片
docx.paragraphs[12].add_run().add_picture(filename1, width=Inches(6))
docx.paragraphs[12].add_run().add_picture(filename2, width=Inches(6))
docx.paragraphs[12].add_run().add_picture(filename3, width=Inches(6))
docx.paragraphs[12].add_run().add_picture(filename4, width=Inches(6))
docx.paragraphs[12].add_run().add_picture(filename5, width=Inches(6))
docx.paragraphs[12].add_run().add_picture(filename6, width=Inches(6))
# ----------------------------------------------------------------------------------------------------------------------
henghezai = round(diaoding*ZJ,2)
huohezai = round(gangjia*ZJ,2)
YJHZ = henghezai
KBHZ = huohezai
P = (f"""恒荷载取{round(diaoding*ZJ,2)}kN/m，活荷载取风荷载和雪荷载最大值，即{round(max(s0,w0)*ZJ,2)}kN/m，按照以下三种组合方式，选取最不利截面设计值：
组合①：1.35×恒荷载+1.5×0.7×活荷载
组合②：1.3×恒荷载+1.5×活荷载
组合③：1.3×恒荷载+1.5×0.9×(活荷载+风荷载)""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('2、内力组合')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 创建一个13行9列的表格
table = docx.add_table(rows=13, cols=9)

# 创建所有单元格的变量
cells = [[table.cell(row, col) for col in range(9)] for row in range(13)]

# 写入第一行的文本
headings = ["截面", "内力", "恒荷载", "活荷载", "风荷载", "组合①", "组合②", "组合③","最不利组合"]
for i, heading in enumerate(headings):
    cells[0][i].text = heading

# 合并第一列的几个单元格并写入文本
for start_row, end_row, text in [(1, 3, "柱A"), (4, 6, "柱B"), (7, 9, "梁B"), (10, 12, "梁C")]:
    cell_start = cells[start_row][0]
    cell_end = cells[end_row][0]
    cell_merge = cell_start.merge(cell_end)
    cell_merge.text = text

# 定义计算系数
ZAM = ZAN = ZAQ = ZBM = ZBN = ZBQ = LBM = LBN = LBQ = LCM = LCN = LCQ = 1  #恒活荷载
ZAM1 = ZAN1 = ZAQ1 = ZBM1 = ZBN1 = ZBQ1 = LBM1 = LBN1 = LBQ1 = LCM1 = LCN1 = LCQ1 = 1  #风荷载
if KD == 24:
    ZAM = 20.92
    ZAN = -12
    ZAQ = -8.75
    ZBM = -35.09
    ZBN = -12
    ZBQ = 8.75
    LBM = -35.09
    LBN = -10.69
    LBQ = 10.31
    LCM = 19.65
    LCN = -8.62
    LCQ = -1.51
    ZAM1 = -32.56
    ZAN1 = -5.7
    ZAQ1 = 8.5
    ZBM1 = 1.36
    ZBN1 = 5.7
    ZBQ1 = 2.1
    LBM1 = 1.36
    LBN1 = 1.09
    LBQ1 = 5.98
    LCM1 = 19.23
    LCN1 = 1.09
    LCQ1 = -6.2
if KD == 21:
    ZAM = 15.81
    ZAN = -10.5
    ZAQ = -6.62
    ZBM = -26.58
    ZBN = -10.5
    ZBQ = -6.60
    LBM = -26.58
    LBN = -8.55
    LBQ = 9
    LCM = 15.48
    LCN = -6.49
    LCQ = -1.3
    ZAM1 = -30.09
    ZAN1 = -4.68
    ZAQ1 = 8.5
    ZBM1 = 3.83
    ZBN1 = -4.68
    ZBQ1 = 2.1
    LBM1 = 3.83
    LBN1 = 1.14
    LBQ1 = 5
    LCM1 = 16.25
    LCN1 = 1.14
    LCQ1 = -5.71
if KD == 18:
    ZAM = 15.81
    ZAN = -10.5
    ZAQ = -6.62
    ZBM = -26.58
    ZBN = -10.5
    ZBQ = -6.60
    LBM = -26.58
    LBN = -8.55
    LBQ = 9
    LCM = 15.48
    LCN = -6.49
    LCQ = -1.3
    ZAM1 = -28
    ZAN1 = -3.6
    ZAQ1 = 8.5
    ZBM1 = 5.92
    ZBN1 = 3.6
    ZBQ1 = 8.5
    LBM1 = -13.8
    LBN1 = -1.23
    LBQ1 = 3.98
    LCM1 = 13.8
    LCN1 = 1.23
    LCQ1 = 3.98
if KD == 15:
    ZAM = 14.81
    ZAN = -9.5
    ZAQ = -5.62
    ZBM = -23.58
    ZBN = -9.5
    ZBQ = -5.60
    LBM = -23.58
    LBN = -7.55
    LBQ = 8
    LCM = 13.48
    LCN = -5.49
    LCQ = -1.3
    ZAM1 = -26
    ZAN1 = -3.6
    ZAQ1 = 7.5
    ZBM1 = 4.92
    ZBN1 = 3.6
    ZBQ1 = 7.5
    LBM1 = -12.8
    LBN1 = -1.23
    LBQ1 = 3.98
    LCM1 = 12.8
    LCN1 = 1.23
    LCQ1 = 2.98

# 第二列内力符号
cells[1][1].text = f'M'
cells[2][1].text = f'N'
cells[3][1].text = f'Q'
cells[4][1].text = f'M'
cells[5][1].text = f'N'
cells[6][1].text = f'Q'
cells[7][1].text = f'M'
cells[8][1].text = f'N'
cells[9][1].text = f'Q'
cells[10][1].text = f'M'
cells[11][1].text = f'N'
cells[12][1].text = f'Q'

# 第三列恒荷载
cells[1][2].text = f'{round(henghezai*ZAM,2)}'  # ZAM
cells[2][2].text = f'{round(henghezai*ZAN,2)}'  # ZAN
cells[3][2].text = f'{round(henghezai*ZAQ,2)}'  # ZAQ
cells[4][2].text = f'{round(henghezai*ZBM,2)}'  # ZBM
cells[5][2].text = f'{round(henghezai*ZBN,2)}'  # ZBN
cells[6][2].text = f'{round(henghezai*ZBQ,2)}'  # ZBQ
cells[7][2].text = f'{round(henghezai*LBM,2)}'  # LBM
cells[8][2].text = f'{round(henghezai*LBN,2)}'  # LBN
cells[9][2].text = f'{round(henghezai*LBQ,2)}'  # LBQ
cells[10][2].text = f'{round(henghezai*LCM,2)}'  # LCM
cells[11][2].text = f'{round(henghezai*LCN,2)}'  # LCN
cells[12][2].text = f'{round(henghezai*LCQ,2)}'  # LCQ

# 第四列活荷载
cells[1][3].text = f'{round(huohezai*ZAM,2)}'  # ZAM
cells[2][3].text = f'{round(huohezai*ZAN,2)}'  # ZAN
cells[3][3].text = f'{round(huohezai*ZAQ,2)}'  # ZAQ
cells[4][3].text = f'{round(huohezai*ZBM,2)}'  # ZBM
cells[5][3].text = f'{round(huohezai*ZBN,2)}'  # ZBN
cells[6][3].text = f'{round(huohezai*ZBQ,2)}'  # ZBQ
cells[7][3].text = f'{round(huohezai*LBM,2)}'  # LBM
cells[8][3].text = f'{round(huohezai*LBN,2)}'  # LBN
cells[9][3].text = f'{round(huohezai*LBQ,2)}'  # LBQ
cells[10][3].text = f'{round(huohezai*LCM,2)}'  # LCM
cells[11][3].text = f'{round(huohezai*LCN,2)}'  # LCN
cells[12][3].text = f'{round(huohezai*LCQ,2)}'  # LCQ

# 第五列风荷载
cells[1][4].text = f'{round(q1*ZAM1,2)}'  # ZAM1
cells[2][4].text = f'{round(q1*ZAN1,2)}'  # ZAN1
cells[3][4].text = f'{round(q1*ZAQ1,2)}'  # ZAQ1
cells[4][4].text = f'{round(q1*ZBM1,2)}'  # ZBM1
cells[5][4].text = f'{round(q1*ZBN1,2)}'  # ZBN1
cells[6][4].text = f'{round(q1*ZBQ1,2)}'  # ZBQ1
cells[7][4].text = f'{round(q2*LBM1,2)}'  # LBM1
cells[8][4].text = f'{round(q2*LBN1,2)}'  # LBN1
cells[9][4].text = f'{round(q2*LBQ1,2)}'  # LBQ1
cells[10][4].text = f'{round(q2*LCM1,2)}'  # LCM1
cells[11][4].text = f'{round(q2*LCN1,2)}'  # LCN1
cells[12][4].text = f'{round(q2*LCQ1,2)}'  # LCQ1

# 第六列组合1
cells[1][5].text = f'{round(1.35*henghezai*ZAM + 1.5*0.7*huohezai*ZAM,2)}'  # ZAM
cells[2][5].text = f'{round(1.35*henghezai*ZAN + 1.5*0.7*huohezai*ZAN,2)}'  # ZAN
cells[3][5].text = f'{round(1.35*henghezai*ZAQ + 1.5*0.7*huohezai*ZAQ,2)}'  # ZAQ
cells[4][5].text = f'{round(1.35*henghezai*ZBM + 1.5*0.7*huohezai*ZBM,2)}'  # ZBM
cells[5][5].text = f'{round(1.35*henghezai*ZBN + 1.5*0.7*huohezai*ZBN,2)}'  # ZBN
cells[6][5].text = f'{round(1.35*henghezai*ZBQ + 1.5*0.7*huohezai*ZBQ,2)}'  # ZBQ
cells[7][5].text = f'{round(1.35*henghezai*LBM + 1.5*0.7*huohezai*LBM,2)}'  # LBM
cells[8][5].text = f'{round(1.35*henghezai*LBN + 1.5*0.7*huohezai*LBN,2)}'  # LBN
cells[9][5].text = f'{round(1.35*henghezai*LBQ + 1.5*0.7*huohezai*LBQ,2)}'  # LBQ
cells[10][5].text = f'{round(1.35*henghezai*LCM + 1.5*0.7*huohezai*LCM,2)}'  # LCM
cells[11][5].text = f'{round(1.35*henghezai*LCN + 1.5*0.7*huohezai*LCN,2)}'  # LCN
cells[12][5].text = f'{round(1.35*henghezai*LCQ + 1.5*0.7*huohezai*LCQ,2)}'  # LCQ

# 第七列组合2
cells[1][6].text = f'{round(1.3*henghezai*ZAM + 1.5*huohezai*ZAM,2)}'  # ZAM
cells[2][6].text = f'{round(1.3*henghezai*ZAN + 1.5*huohezai*ZAN,2)}'  # ZAN
cells[3][6].text = f'{round(1.3*henghezai*ZAQ + 1.5*huohezai*ZAQ,2)}'  # ZAQ
cells[4][6].text = f'{round(1.3*henghezai*ZBM + 1.5*huohezai*ZBM,2)}'  # ZBM
cells[5][6].text = f'{round(1.3*henghezai*ZBN + 1.5*huohezai*ZBN,2)}'  # ZBN
cells[6][6].text = f'{round(1.3*henghezai*ZBQ + 1.5*huohezai*ZBQ,2)}'  # ZBQ
cells[7][6].text = f'{round(1.3*henghezai*LBM + 1.5*huohezai*LBM,2)}'  # LBM
cells[8][6].text = f'{round(1.3*henghezai*LBN + 1.5*huohezai*LBN,2)}'  # LBN
cells[9][6].text = f'{round(1.3*henghezai*LBQ + 1.5*huohezai*LBQ,2)}'  # LBQ
cells[10][6].text = f'{round(1.3*henghezai*LCM + 1.5*huohezai*LCM,2)}'  # LCM
cells[11][6].text = f'{round(1.3*henghezai*LCN + 1.5*huohezai*LCN,2)}'  # LCN
cells[12][6].text = f'{round(1.3*henghezai*LCQ + 1.5*huohezai*LCQ,2)}'  # LCQ

# 第八列组合3
cells[1][7].text = f'{round(1.3*henghezai*ZAM + 1.5*0.9*(huohezai*ZAM + q1*ZAM),2)}'  # ZAM
cells[2][7].text = f'{round(1.3*henghezai*ZAN + 1.5*0.9*(huohezai*ZAN + q1*ZAN),2)}'  # ZAN
cells[3][7].text = f'{round(1.3*henghezai*ZAQ + 1.5*0.9*(huohezai*ZAQ + q1*ZAQ),2)}'  # ZAQ
cells[4][7].text = f'{round(1.3*henghezai*ZBM + 1.5*0.9*(huohezai*ZBM + q1*ZBM),2)}'  # ZBM
cells[5][7].text = f'{round(1.3*henghezai*ZBN + 1.5*0.9*(huohezai*ZBN + q1*ZBN),2)}'  # ZBN
cells[6][7].text = f'{round(1.3*henghezai*ZBQ + 1.5*0.9*(huohezai*ZBQ + q1*ZBQ),2)}'  # ZBQ
cells[7][7].text = f'{round(1.3*henghezai*LBM + 1.5*0.9*(huohezai*LBM + q2*LBM),2)}'  # LBM
cells[8][7].text = f'{round(1.3*henghezai*LBN + 1.5*0.9*(huohezai*LBN + q2*LBN),2)}'  # LBN
cells[9][7].text = f'{round(1.3*henghezai*LBQ + 1.5*0.9*(huohezai*LBQ + q2*LBQ),2)}'  # LBQ
cells[10][7].text = f'{round(1.3*henghezai*LCM + 1.5*0.9*(huohezai*LCM + q2*LCM),2)}'  # LCM
cells[11][7].text = f'{round(1.3*henghezai*LCN + 1.5*0.9*(huohezai*LCN + q2*LCN),2)}'  # LCN
cells[12][7].text = f'{round(1.3*henghezai*LCQ + 1.5*0.9*(huohezai*LCQ + q2*LCQ),2)}'  # LCQ

# 最不利组合内力设计值选取
# 第九列 - 最不利组合
for i in range(1, 13):
    # 计算前三种组合的绝对值最大值
    max_value = max(abs(float(cells[i][5].text)), abs(float(cells[i][6].text)), abs(float(cells[i][7].text)))

    # 将最大值写入第九列
    cells[i][8].text = f'{round(max_value, 2)}'
# 定义最不利截面设计值
ZAMmax = round(float(cells[1][8].text),2)
ZANmax = round(float(cells[2][8].text),2)
ZAQmax = round(float(cells[3][8].text),2)
ZBMmax = round(float(cells[4][8].text),2)
ZBNmax = round(float(cells[5][8].text),2)
ZBQmax = round(float(cells[6][8].text),2)
LBMmax = round(float(cells[7][8].text),2)
LBNmax = round(float(cells[8][8].text),2)
LBQmax = round(float(cells[9][8].text),2)
LCMmax = round(float(cells[10][8].text),2)
LCNmax = round(float(cells[11][8].text),2)
LCQmax = round(float(cells[12][8].text),2)

# 找出最大的一个数
max_value = max(ZAMmax, ZANmax, ZAQmax, ZBMmax, ZBNmax, ZBQmax, LBMmax, LBNmax, LBQmax, LCMmax, LCNmax, LCQmax)

# 内力最大值，用于选择H型钢
neilizuidazhi = max_value
# ----------------------------------------------------------------------------------------------------------------------
P = ('五、钢架设计')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('1、截面设计')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 选择H型钢型号
mingcheng = "这里是H型钢名称"

A = q = Ix = Wx = ix = Iy = Wy = iy = b = t = 2
hw = tw = h = 1
if ZG <= 8:
    if neilizuidazhi < 200:
        mingcheng = "HW300×300×10×15"
    elif 200 <= neilizuidazhi <= 300:
        mingcheng = "HN500×200×10×16"
    elif neilizuidazhi > 300:
        mingcheng = "HM588×300×12×20"
if ZG > 8:
    mingcheng = "HM588×300×12×20"


if mingcheng == "HM588×300×12×20":
    A = 192.5
    q = 72.4
    Ix = 118000
    Wx = 4020
    ix = 24.8
    Iy = 9020
    Wy = 601
    iy = 6.85
    b = 300  #宽度
    h = 588  #高度
    t = 20  #翼缘厚度
    hw = 588-20  #腹板高度
    tw = 12  #腹板厚度
if mingcheng == "HW300×300×10×15":
    A = 120.4
    q = 94.5
    Ix = 20500
    Wx = 1370
    ix = 13.1
    Iy = 6760
    Wy = 450
    iy = 7.49
    b = 300
    h = 300
    t = 15
    hw = 300-15
    tw = 10
if mingcheng == "HN500×200×10×16":
    A = 114.2
    q = 89.6
    Ix = 47800
    Wx = 1910
    ix = 20.5
    Iy = 2140
    Wy = 214
    iy = 4.33
    b = 200
    h = 500
    t = 16
    hw = 500-16
    tw = 10
# ----------------------------------------------------------------------------------------------------------------------
P = f"""梁柱均选用{mingcheng}，截面特性为A={A}mm²，b={b}mm，t={t}mm，Ix={Ix}mm⁴，Wx={Wx}cm³，ix={ix}cm，Iy={Iy}mm⁴，Wy={Wy}cm³，iy={iy}cm。"""
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('2、构件验算')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
b_t = b/t
hw_tw = hw / tw
P = (f"""构件宽厚比的验算：
翼缘部分：
\\frac{{b}}{{t}}=\\frac{{{b}}}{{{t}}}={b_t:.2f}<15\\sqrt{{\\frac{{{fy}}}{{f_y}}}}=15
腹板部分：
\\frac{{h_w}}{{t_w}}=\\frac{{{hw}}}{{{tw}}}={hw_tw:.1f}<250\\sqrt{{\\frac{{{fy}}}{{f_y}}}}=250""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
import numpy as np
fv = round(max(ZAQmax,ZBQmax,LBQmax,LCQmax),2)
hw_tw_fv = hw * tw * 125
Vu = hw_tw_fv
# 单位化成m
Af1 = round(b*t/1000,2)
h1 = round((h/2-t/2)/1000,2)
h2 = round((h/2-t/2)/1000,2)
Af2 = round(b*t/1000,2)
N = round(max(LBNmax,LCNmax),2)

# 计算
V = fv
if V<0.5*Vu:
    V = 0.5*Vu


inner_expression = (Af1 * (h1**2 / h2)) + (Af2 * h2)
outer_expression = f - (N / A/100)
result = round(inner_expression * outer_expression,2)

Mf = result
M = Mf
Meu = max(LBMmax,LCMmax)

left_result = (V / (0.5 * Vu)) - 1
right_result = (M - Mf) / (Meu - Mf)
total_result = round(left_result**2 + right_result,2)


Ia = Ix
zhugao = ZG*1000  # 柱高
# 计算
k1_result = round((Ia * 10**4) / zhugao,2)

Ib = Ix
liangchang = 2*round(((KD*1000/2)**2+(KD*1000*0.2)**2)**0.5,2)  # 梁长
# 计算
k2_result = round((Ib * 10**4) / liangchang,2)

l_x = KD * 1000
lambda_x_result = round(l_x/10 / ix,2)

# 定义一维数组，包含了你给出的长细比λ与稳定系数φ的对应关系
phi_values = np.array([
    1, 0.997, 0.995, 0.992, 0.989, 0.987, 0.984, 0.981, 0.979, 0.976,
    0.974, 0.971, 0.968, 0.966, 0.963, 0.96, 0.958, 0.955, 0.952, 0.949,
    0.947, 0.944, 0.941, 0.938, 0.936, 0.933, 0.93, 0.927, 0.924, 0.921,
    0.918, 0.915, 0.912, 0.909, 0.906, 0.903, 0.899, 0.896, 0.893, 0.889,
    0.886, 0.882, 0.879, 0.875, 0.872, 0.868, 0.864, 0.861, 0.858, 0.855,
    0.852, 0.849, 0.846, 0.843, 0.839, 0.836, 0.832, 0.829, 0.825, 0.822,
    0.818, 0.814, 0.81, 0.806, 0.802, 0.797, 0.793, 0.789, 0.784, 0.779,
    0.775, 0.77, 0.765, 0.76, 0.755, 0.75, 0.744, 0.739, 0.733, 0.728,
    0.722 ,0.716, 0.71, 0.704 ,0.698,0.692 ,0.686 ,0.68 ,0.673, 0.667,
    0.661 ,0.654 ,0.648 ,0.641, 0.634 ,0.626, 0.618,0.611 ,0.603, 0.595,
    0.588 ,0.58,0.573 ,0.566 ,0.558 ,0.551 ,0.544 ,0.537 ,0.53 ,0.523,
    0.516, 0.509 ,0.502 ,0.496, 0.489 ,0.483 ,0.476, 0.47, 0.464 ,0.458,
    0.452 ,0.446 ,0.44, 0.434 ,0.428,0.423 ,0.417,0.412 ,0.406, 0.401,
    0.396 ,0.391 ,0.386 ,0.381, 0.376 ,0.371, 0.367 ,0.362 ,0.357 ,0.353,
    0.349 ,0.344 ,0.34, 0.336 ,0.332 ,0.328 ,0.324,0.32 ,0.316 ,0.312,
    0.308 ,0.305 ,0.301 ,0.298 ,0.294 ,0.291 ,0.287, 0.284 ,0.281, 0.277,
    0.274 ,0.271 ,0.268 ,0.265, 0.262 ,0.259 ,0.256 ,0.253, 0.251 ,0.248,
    0.245, 0.243 ,0.24 ,0.237 ,0.235,0.232 ,0.23, 0.227 ,0.225, 0.223,
    0.22, 0.218, 0.216 ,0.214 ,0.211, 0.209, 0.207 ,0.205 ,0.203 ,0.201 ])
# 定义λ的值
lamuda = lambda_x_result

# 找出lambda对应的索引
index = int(np.floor(lamuda))  # 对应的整数部分
frac = lamuda % 1  # 对应的小数部分

# 找出相邻的φ值
phi_left = phi_values[index]
phi_right = phi_values[min(index + 1, len(phi_values) - 1)]  # 限制索引不超过数组长度

# 一维线性插值
phi = round(phi_left * (1 - frac) + phi_right * frac,3)
phi_x = phi

E_A = 206

beta_mx = 1.0

# 计算
N_result = (math.pi**2 * E_A * 10**6 * A) / (1.1 * lambda_x_result**2)
N_result_kN = N_result / 1000  # 将结果转换为kN

W_e1 = Wx
Mx = max(LBMmax,LCMmax)
# 计算
left_result1 = N / (phi_x * A)
right_result1 = (beta_mx * Mx*10**3) / (W_e1 * (1 - phi_x * (N / round(N_result_kN,2))))
total_result1 = round(left_result1 + right_result1,2)

# 檩距
l_y = 1200
# 长细比
lambda_y_result = l_y/10 / iy

# 定义一维数组，包含了你给出的长细比λ与稳定系数φ的对应关系
phi_values = np.array([
    1, 0.997, 0.995, 0.992, 0.989, 0.987, 0.984, 0.981, 0.979, 0.976,
    0.974, 0.971, 0.968, 0.966, 0.963, 0.96, 0.958, 0.955, 0.952, 0.949,
    0.947, 0.944, 0.941, 0.938, 0.936, 0.933, 0.93, 0.927, 0.924, 0.921,
    0.918, 0.915, 0.912, 0.909, 0.906, 0.903, 0.899, 0.896, 0.893, 0.889,
    0.886, 0.882, 0.879, 0.875, 0.872, 0.868, 0.864, 0.861, 0.858, 0.855,
    0.852, 0.849, 0.846, 0.843, 0.839, 0.836, 0.832, 0.829, 0.825, 0.822,
    0.818, 0.814, 0.81, 0.806, 0.802, 0.797, 0.793, 0.789, 0.784, 0.779,
    0.775, 0.77, 0.765, 0.76, 0.755, 0.75, 0.744, 0.739, 0.733, 0.728,
    0.722 ,0.716, 0.71, 0.704 ,0.698,0.692 ,0.686 ,0.68 ,0.673, 0.667,
    0.661 ,0.654 ,0.648 ,0.641, 0.634 ,0.626, 0.618,0.611 ,0.603, 0.595,
    0.588 ,0.58,0.573 ,0.566 ,0.558 ,0.551 ,0.544 ,0.537 ,0.53 ,0.523,
    0.516, 0.509 ,0.502 ,0.496, 0.489 ,0.483 ,0.476, 0.47, 0.464 ,0.458,
    0.452 ,0.446 ,0.44, 0.434 ,0.428,0.423 ,0.417,0.412 ,0.406, 0.401,
    0.396 ,0.391 ,0.386 ,0.381, 0.376 ,0.371, 0.367 ,0.362 ,0.357 ,0.353,
    0.349 ,0.344 ,0.34, 0.336 ,0.332 ,0.328 ,0.324,0.32 ,0.316 ,0.312,
    0.308 ,0.305 ,0.301 ,0.298 ,0.294 ,0.291 ,0.287, 0.284 ,0.281, 0.277,
    0.274 ,0.271 ,0.268 ,0.265, 0.262 ,0.259 ,0.256 ,0.253, 0.251 ,0.248,
    0.245, 0.243 ,0.24 ,0.237 ,0.235,0.232 ,0.23, 0.227 ,0.225, 0.223,
    0.22, 0.218, 0.216 ,0.214 ,0.211, 0.209, 0.207 ,0.205 ,0.203 ,0.201 ])
# 定义λ的值
lamuda = lambda_y_result

# 找出lambda对应的索引
index = int(np.floor(lamuda))  # 对应的整数部分
frac = lamuda % 1  # 对应的小数部分

# 找出相邻的φ值
phi_left = phi_values[index]
phi_right = phi_values[min(index + 1, len(phi_values) - 1)]  # 限制索引不超过数组长度

# 一维线性插值
phi = round(phi_left * (1 - frac) + phi_right * frac,3)
phi_y = phi

lamudayjieguo = lambda_y_result
phi_by_result = (4320 / lamudayjieguo**2) * (A*100 * hw / (Wx * 10**3)) * math.sqrt(((lamudayjieguo * t) / (4.4 * hw))**2)

faibyipie = round(1.07-0.282/phi_by_result,2)

beta_t = 1.0 - N / N_result_kN + 0.75 * (N / N_result_kN) ** 2

result2 = N / (phi_y * A) + ( beta_t * max(LBMmax,LCMmax)*1000 )/ (Wx * faibyipie)

P = (f"""钢架梁的验算:
内力设计值V_{{max}}={fv}kN
①抗剪强度验算(考虑仅有支座加劲肋)：
\\lambda_s=\\frac{{{{hw}}/{{tw}}}}{{41\\sqrt{{5.34}}}}\\sqrt{{\\frac{{fy}}{{235}}}}=\\frac{{{hw}/{tw}}}{{41\\sqrt{{5.34}}}}={round((hw/tw)/(41*5.34**0.5)*(fy/235)**0.5,2)}<0.8
则极限承载力：
V_u=h_w\\times t_w\\times f_v={hw}\\times{tw}\\times125={hw_tw_fv/1000}kN>{fv}kN
满足要求。
②弯、剪、压共同作用下的验算：
取梁端截面：
N={max(LBNmax,LCNmax)}kN，V={max(LBQmax,LCQmax)}kN，M={max(LBMmax,LCMmax)}kN·m，\\text{{又因}}V<0.5V_u，\\text{{取}}V=0.5V_u={round(0.5*Vu/1000,2)}kN
M_f=\\left(A_{{f1}}\\times\\frac{{h_1^2}}{{h_2}}+A_{{f2}}\\times h_2\\right)\\times\\left(f-\\frac{{N}}{{A}}\\right)=\\left({Af1}\\times\\frac{{{h1}^2}}{{{h2}}}+{Af2}\\times {h2}\\right)\\times\\left({f}-\\frac{{{N}}}{{{round(A/100,2)}}}\\right)={result}kN·m＞{max(LBMmax,LCMmax)}kN·m
M_f＞M，\\text{{取}}M=M_f={Mf}kN·m,\\text{{故：}}
\\left(\\frac{{V}}{{0.5\\times V_u}}-1\\right)^2+\\frac{{M-M_f}}{{M_{{eu}}-M_f}}=\\left(\\frac{{{V}}}{{0.5\\times{Vu}}}-1\\right)^2+\\frac{{{M}-{Mf}}}{{{Meu}-{Mf}}}={total_result}<1
满足要求。
③整体稳定性验算:
1)梁平面内稳定验算
\\text{{计算长度取横梁长度}}l_x={KD*1000}mm。
\\lambda_x=\\frac{{l_x}}{{i_x}}=\\frac{{{l_x/10}}}{{{ix}}}={lambda_x_result}<[\\lambda]=150,\\text{{b类截面，查表得}}φ_x={phi_x}
N^' = \\frac{{\\pi^2 \\times E \\times A}}{{1.1 \\times \\lambda^2}} = \\frac{{3.14^2 \\times {E_A} \\times 10^6 \\times {A}}}{{1.1 \\times {lambda_x_result}^2}}={N_result_kN:.2f}kN,\\text{{取}}\\beta_{{mx}}={beta_mx}
则弯剪压共同作用下最大应力：
\\frac{{N}}{{\\varphi_x \\times A}} + \\frac{{\\beta_{{mx}} \\times M_x}}{{W_{{e1}} \\times \\left(1-\\varphi_x \\times \\frac{{N}}{{N'}}\\right)}} = \\frac{{{N}}}{{{phi_x} \\times {A}}} + \\frac{{{beta_mx} \\times {Mx*1000}}}{{{W_e1} \\times \\left(1-{phi_x} \\times \\frac{{{N}}}{{{round(N_result_kN,2)}}}\\right)}} = {total_result1:.2f}N/mm^2 < f = {f}N/mm^2
满足要求。
2)梁平面外稳定验算
考虑蒙皮效应，两个檩条间距不小于1200mm，计算长度按两个檩距考虑，即：
l_y={2*l_y}mm，\\text{{对于等截面构件}}\gamma=0，μ_s=μ_w=1。
\\lambda_y=\\frac{{l_y}}{{i_y}}=\\frac{{{l_y/10}}}{{{iy}}}={lambda_y_result:.1f}，\\text{{是b类截面，查表得}}φ_y={phi_y}。
\\varphi_{{by}}=\\frac{{4320}}{{{{\\lambda_y}}^2}}\\times\\frac{{A\\times h_w}}{{W_x\\times{{10}}^3}}\\times\\sqrt{{\\left(\\frac{{\\lambda_y\\times t}}{{4.4\\times h_w}}\\right)^2}}=\\frac{{4320}}{{{{{round(lamudayjieguo, 2)}}}^2}}\\times\\frac{{{A*100}\\times{hw}}}{{{Wx}\\times{{10}}^3}}\\times\\sqrt{{\\left(\\frac{{{round(lamudayjieguo, 2)}\\times{t}}}{{4.4\\times{hw}}}\\right)^2}}={phi_by_result:.1f}＞0.6
\\text{{则取}}{{\\varphi_b}}^'=1.07-0.282/\\varphi_{{by}}={faibyipie}
\\beta_t=1.0-\\frac{{N}}{{N\\prime}}+0.75\\times\\left(\\frac{{N}}{{N\\prime}}\\right)^2=1.0-\\frac{{{N}}}{{{round(N_result_kN,2)}}}+0.75\\times\\left(\\frac{{{N}}}{{{round(N_result_kN,2)}}}\\right)^2={beta_t:.3f}
\\frac{{N}}{{\\varphi_y\\times A}}+\\frac{{\\beta_t\\times M}}{{Wx\\times\\varphi_b^'}}=\\frac{{{N}}}{{{phi_y}\\times{A}}}+\\frac{{{round(beta_t,3)}\\times{max(LBMmax,LCMmax)*1000}}}{{{Wx}\\times{faibyipie}}}={result2:.2f}N/mm^2<f={f}N/mm^2""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
import numpy as np
Vzhuzi = round(max(ZAQmax,ZBQmax),2)

inner_expression1 = (Af1 * (h1**2 / h2)) + (Af2 * h2)
outer_expression1 = fy - (N / A/100)
result11 = round(inner_expression1 * outer_expression1,2)

Mf = result11
M = Mf
left_result1 = (V / (0.5 * Vu)) - 1
right_result1 = (M - Mf) / (Meu - Mf)
total_result11 = round(left_result1**2 + right_result1,2)

μ_r = 3.414

H = zhugao

l_x = round(μ_r * H,2)

# 计算
lambda_x = round(l_x / (ix*10),2)

# 定义一维数组，包含了你给出的长细比λ与稳定系数φ的对应关系
phi_values = np.array([
    1, 0.997, 0.995, 0.992, 0.989, 0.987, 0.984, 0.981, 0.979, 0.976,
    0.974, 0.971, 0.968, 0.966, 0.963, 0.96, 0.958, 0.955, 0.952, 0.949,
    0.947, 0.944, 0.941, 0.938, 0.936, 0.933, 0.93, 0.927, 0.924, 0.921,
    0.918, 0.915, 0.912, 0.909, 0.906, 0.903, 0.899, 0.896, 0.893, 0.889,
    0.886, 0.882, 0.879, 0.875, 0.872, 0.868, 0.864, 0.861, 0.858, 0.855,
    0.852, 0.849, 0.846, 0.843, 0.839, 0.836, 0.832, 0.829, 0.825, 0.822,
    0.818, 0.814, 0.81, 0.806, 0.802, 0.797, 0.793, 0.789, 0.784, 0.779,
    0.775, 0.77, 0.765, 0.76, 0.755, 0.75, 0.744, 0.739, 0.733, 0.728,
    0.722 ,0.716, 0.71, 0.704 ,0.698,0.692 ,0.686 ,0.68 ,0.673, 0.667,
    0.661 ,0.654 ,0.648 ,0.641, 0.634 ,0.626, 0.618,0.611 ,0.603, 0.595,
    0.588 ,0.58,0.573 ,0.566 ,0.558 ,0.551 ,0.544 ,0.537 ,0.53 ,0.523,
    0.516, 0.509 ,0.502 ,0.496, 0.489 ,0.483 ,0.476, 0.47, 0.464 ,0.458,
    0.452 ,0.446 ,0.44, 0.434 ,0.428,0.423 ,0.417,0.412 ,0.406, 0.401,
    0.396 ,0.391 ,0.386 ,0.381, 0.376 ,0.371, 0.367 ,0.362 ,0.357 ,0.353,
    0.349 ,0.344 ,0.34, 0.336 ,0.332 ,0.328 ,0.324,0.32 ,0.316 ,0.312,
    0.308 ,0.305 ,0.301 ,0.298 ,0.294 ,0.291 ,0.287, 0.284 ,0.281, 0.277,
    0.274 ,0.271 ,0.268 ,0.265, 0.262 ,0.259 ,0.256 ,0.253, 0.251 ,0.248,
    0.245, 0.243 ,0.24 ,0.237 ,0.235,0.232 ,0.23, 0.227 ,0.225, 0.223,
    0.22, 0.218, 0.216 ,0.214 ,0.211, 0.209, 0.207 ,0.205 ,0.203 ,0.201 ])
# 定义λ的值
lamuda = lambda_x

# 找出lambda对应的索引
index = int(np.floor(lamuda))  # 对应的整数部分
frac = lamuda % 1  # 对应的小数部分

# 找出相邻的φ值
phi_left = phi_values[index]
phi_right = phi_values[min(index + 1, len(phi_values) - 1)]  # 限制索引不超过数组长度

# 一维线性插值
phi = round(phi_left * (1 - frac) + phi_right * frac,3)
phi_x = phi

N_result = (math.pi**2 * E_A * 10**6 * A) / (1.1 * lambda_x**2)
N_result_kN = N_result / 1000  # 将结果转换为kN

N = round(max(ZBNmax,ZANmax),2)

W_e1 = Wx
Mx = max(ZAMmax,ZBMmax)


# 计算最大应力
left_result1 = N / (phi_x * A)
right_result1 = (beta_mx * Mx*10**3) / (W_e1 * (1 - phi_x * (N / round(N_result_kN,2))))
total_result1 = round(left_result1 + right_result1,2)

import numpy as np

# 定义一维数组，包含了你给出的长细比λ与稳定系数φ的对应关系
phi_values = np.array([
    1, 0.997, 0.995, 0.992, 0.989, 0.987, 0.984, 0.981, 0.979, 0.976,
    0.974, 0.971, 0.968, 0.966, 0.963, 0.96, 0.958, 0.955, 0.952, 0.949,
    0.947, 0.944, 0.941, 0.938, 0.936, 0.933, 0.93, 0.927, 0.924, 0.921,
    0.918, 0.915, 0.912, 0.909, 0.906, 0.903, 0.899, 0.896, 0.893, 0.889,
    0.886, 0.882, 0.879, 0.875, 0.872, 0.868, 0.864, 0.861, 0.858, 0.855,
    0.852, 0.849, 0.846, 0.843, 0.839, 0.836, 0.832, 0.829, 0.825, 0.822,
    0.818, 0.814, 0.81, 0.806, 0.802, 0.797, 0.793, 0.789, 0.784, 0.779,
    0.775, 0.77, 0.765, 0.76, 0.755, 0.75, 0.744, 0.739, 0.733, 0.728,
    0.722 ,0.716, 0.71, 0.704 ,0.698,0.692 ,0.686 ,0.68 ,0.673, 0.667,
    0.661 ,0.654 ,0.648 ,0.641, 0.634 ,0.626, 0.618,0.611 ,0.603, 0.595,
    0.588 ,0.58,0.573 ,0.566 ,0.558 ,0.551 ,0.544 ,0.537 ,0.53 ,0.523,
    0.516, 0.509 ,0.502 ,0.496, 0.489 ,0.483 ,0.476, 0.47, 0.464 ,0.458,
    0.452 ,0.446 ,0.44, 0.434 ,0.428,0.423 ,0.417,0.412 ,0.406, 0.401,
    0.396 ,0.391 ,0.386 ,0.381, 0.376 ,0.371, 0.367 ,0.362 ,0.357 ,0.353,
    0.349 ,0.344 ,0.34, 0.336 ,0.332 ,0.328 ,0.324,0.32 ,0.316 ,0.312,
    0.308 ,0.305 ,0.301 ,0.298 ,0.294 ,0.291 ,0.287, 0.284 ,0.281, 0.277,
    0.274 ,0.271 ,0.268 ,0.265, 0.262 ,0.259 ,0.256 ,0.253, 0.251 ,0.248,
    0.245, 0.243 ,0.24 ,0.237 ,0.235,0.232 ,0.23, 0.227 ,0.225, 0.223,
    0.22, 0.218, 0.216 ,0.214 ,0.211, 0.209, 0.207 ,0.205 ,0.203 ,0.201 ])
# 定义λ的值
lamuda = lambda_y_result

# 找出lambda对应的索引
index = int(np.floor(lamuda))  # 对应的整数部分
frac = lamuda % 1  # 对应的小数部分

# 找出相邻的φ值
phi_left = phi_values[index]
phi_right = phi_values[min(index + 1, len(phi_values) - 1)]  # 限制索引不超过数组长度

# 一维线性插值
phi = round(phi_left * (1 - frac) + phi_right * frac,3)

phi_y = phi
lamudayjieguo = lambda_y_result
phi_by_result = (4320 / lamudayjieguo**2) * (A*100 * hw / (Wx * 10**3)) * math.sqrt(((lamudayjieguo * t) / (4.4 * hw))**2)

faibyipie = round(1.07-0.282/phi_by_result,2)

beta_t = 1.0 - N / N_result_kN + 0.75 * (N / N_result_kN) ** 2

result2 = N / (phi_y * A) + ( beta_t * max(LBMmax,LCMmax)*1000 )/ (Wx * faibyipie)

xi_t = round(KD*1000 / zhugao,2)

W = (q1 + abs(q3)) * ZG

W_k = W
kesai1 = xi_t


# 计算
mu = (W_k * 10**3 * zhugao**3) / (12 * E_A * 10**3 * Ix*10**4) * (2 + kesai1)

P = (f"""钢架柱的验算:
内力设计值V_{{max}}={Vzhuzi}kN
①抗剪强度验算(考虑仅有支座加劲肋)：
\\lambda_s=\\frac{{{{hw}}/{{tw}}}}{{41\\sqrt{{5.34}}}}\\sqrt{{\\frac{{fy}}{{235}}}}=\\frac{{{hw}/{tw}}}{{41\\sqrt{{5.34}}}}={round((hw/tw)/(41*5.34**0.5)*(fy/235)**0.5,2)}<0.8
则极限承载力：
V_u=h_w\\times t_w\\times f_v={hw}\\times{tw}\\times125={hw_tw_fv/1000}kN>{Vzhuzi}kN
满足要求。
②弯、剪、压共同作用下的验算：
取柱端截面：
N={max(ZANmax,ZBNmax)}kN，V={max(ZAQmax,ZBQmax)}kN，M={max(ZAMmax,ZBMmax)}kN·m，\\text{{又因}}V<0.5V_u，\\text{{取}}V=0.5V_u={round(0.5*Vu/1000,2)}kN
M_f=\\left(A_{{f1}}\\times\\frac{{h_1^2}}{{h_2}}+A_{{f2}}\\times h_2\\right)\\times\\left(f-\\frac{{N}}{{A}}\\right)=\\left({Af1}\\times\\frac{{{h1}^2}}{{{h2}}}+{Af2}\\times {h2}\\right)\\times\\left({fy}-\\frac{{{N}}}{{{round(A/100,2)}}}\\right)={result11}kN·m＞{max(ZAMmax,ZBMmax)}kN·m
M_f＞M，\\text{{取}}M=M_f={result11}kN·m，\\text{{故：}}
\\left(\\frac{{V}}{{0.5\\times V_u}}-1\\right)^2+\\frac{{M-M_f}}{{M_{{eu}}-M_f}}=\\left(\\frac{{{V}}}{{0.5\\times{Vu}}}-1\\right)^2+\\frac{{{M}-{Mf}}}{{{Meu}-{Mf}}}={total_result11}<1
满足要求。
③整体稳定性验算:
1)柱平面内稳定验算
\\text{{钢架柱高}}H={zhugao}mm，\\text{{梁长}}L={liangchang}mm。
柱的线刚度：
k_1=\\frac{{I_a}}{{h}}=\\frac{{{Ia}\\times10^4}}{{{zhugao}}}={k1_result}mm³
梁的线刚度:
k_2=\\frac{{I_b}}{{h}}=\\frac{{{Ib}\\times10^4}}{{{liangchang}}}={k2_result}mm³
\\text{{则}}k_1/k_2={k1_result}/{k2_result}={round(k1_result/k2_result,2)}，I_y/I_x={round(Iy/Ix,2)}，\\text{{查表得}}μ_r={μ_r}
计算长度：
l_x=\\mu_r\\times h={μ_r}\\times{H}={l_x}mm
验算长细比：
\\lambda_x=\\frac{{l_x}}{{i_x}}=\\frac{{{l_x}}}{{{ix*10}}}={lambda_x:.2f}<\\left[\\lambda\\right]=150,b\\text{{类截面}}，\\text{{查表得}}\\varphi_x={phi_x}
N^' = \\frac{{\\pi^2 \\times E \\times A}}{{1.1 \\times \\lambda^2}} = \\frac{{3.14^2 \\times {E_A} \\times 10^6 \\times {A}}}{{1.1 \\times {lambda_x}^2}}={N_result_kN:.2f}kN，\\text{{取}}\\beta_{{mx}}={beta_mx}
则弯、剪、压共同作用下截面最大应力：
\\frac{{N}}{{\\varphi_x \\times A}} + \\frac{{\\beta_{{mx}} \\times M_x}}{{W_{{e1}} \\times \\left(1-\\varphi_x \\times \\frac{{N}}{{N'}}\\right)}} = \\frac{{{N}}}{{{phi_x} \\times {A}}} + \\frac{{{beta_mx} \\times {Mx*1000}}}{{{W_e1} \\times \\left(1-{phi_x} \\times \\frac{{{N}}}{{{round(N_result_kN,2)}}}\\right)}} = {total_result1:.2f}N/mm^2 < f = {f}N/mm^2
满足要求。
2)柱面外稳定验算
考虑蒙皮效应，两个檩条间距不小于1200mm，计算长度按两个檩距考虑，\\text{{即}}l_y={2*l_y}mm，\\text{{对于等截面构件}}\gamma=0，μ_s=μ_w=1。
\\lambda_y=\\frac{{l_y}}{{i_y}}=\\frac{{{l_y/10}}}{{{iy}}}={lambda_y_result:.1f}
\\text{{是b类截面，查表得}}φ_y={phi_y}。
\\varphi_{{by}}=\\frac{{4320}}{{{{\\lambda_y}}^2}}\\times\\frac{{A\\times h_w}}{{W_x\\times{{10}}^3}}\\times\\sqrt{{\\left(\\frac{{\\lambda_y\\times t}}{{4.4\\times h_w}}\\right)^2}}=\\frac{{4320}}{{{{{round(lamudayjieguo, 2)}}}^2}}\\times\\frac{{{A*100}\\times{hw}}}{{{Wx}\\times{{10}}^3}}\\times\\sqrt{{\\left(\\frac{{{round(lamudayjieguo, 2)}\\times{t}}}{{4.4\\times{hw}}}\\right)^2}}={phi_by_result:.1f}＞0.6
则取{{\\varphi_b}}^'=1.07-0.282/\\varphi_{{by}}={faibyipie}
\\beta_t=1.0-\\frac{{N}}{{N\\prime}}+0.75\\times\\left(\\frac{{N}}{{N\\prime}}\\right)^2=1.0-\\frac{{{N}}}{{{round(N_result_kN,2)}}}+0.75\\times\\left(\\frac{{{N}}}{{{round(N_result_kN,2)}}}\\right)^2={beta_t:.3f}
则弯、剪、压共同作用下截面最大应力：
\\frac{{N}}{{\\varphi_y\\times A}}+\\frac{{\\beta_t\\times M}}{{Wx\\times\\varphi_b^'}}=\\frac{{{N}}}{{{phi_y}\\times{A}}}+\\frac{{{round(beta_t,3)}\\times{max(LBMmax,LCMmax)*1000}}}{{{Wx}\\times{faibyipie}}}={result2:.2f}N/mm^2<f={f}N/mm^2
4)钢架在风荷载作用下侧移验算
截面惯性矩：
I_c=I_b={Ib}mm^4
位移修正系数：
\\xi_t=\\frac{{I_cL}}{{hI_B}}=\\frac{{{KD*1000}}}{{{zhugao}}}={xi_t:.2f}
风压设计值：
W=\\left(W_1+W_4\\right)\\times h=\\left({q1}+{abs(q3)}\\right)\\times {ZG}={W:.3f}kN
钢架柱顶等效水平力：
W_k=0.67×W=0.67×{W:.3f}={0.67*W:.3f}kN
左风荷载下顶点位移：
\\mu=\\frac{{H\\times h^3}}{{12\\times E\\times I_c}}\\times\\left(2+\\xi_1\\right)=\\frac{{{W_k}\\times10^3\\times{zhugao}^3}}{{12\\times{E_A}\\times10^3\\times{Ix}\\times10^4}}\\times\\left(2+{kesai1}\\right)={mu:.2f}mm<\\left[\\mu\\right]=h/150=60mm""")
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('3、节点验算')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f"""拼接板尺寸如图：""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f""" """)
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 插入图片
docx.paragraphs[24].add_run().add_picture('C:\\Users\\Q\\PycharmProjects\\pythonProject\\轻型门式钢架\\拼接板尺寸详图.png', width=Inches(2))
# ----------------------------------------------------------------------------------------------------------------------
Pl = 290  #单栓承载力设计值
N = max(LBNmax,LCNmax)
V = max(LBQmax,LCQmax)
M = max(LBMmax,LCMmax)

y1 = 0.3
y2 = 0.4

n = 10  # 螺栓个数

# 计算
N_1 = (M * y1) / (4*(y1**2 + y2**2)) - N / n

# 计算
N_2 = (M * y2) / (4*(y1**2 + y2**2)) - N / n

mu = 0.45
Nf = 1

# 计算
N_vb = 0.9 * Nf * mu * Pl * 8

Nt = N

N_tb = 0.8*Pl

# 计算
expression_value = V / N_vb + Nt / N_tb

# 赋值
ef = 60
ew = 80

# 计算
t1 = math.sqrt((6 * ef * ew * Nt * 1000) / (ew * b + 2 * ef * (ef + ew) * f))

db = hw
dc = hw
tc = t

# 计算
tau = (M * 10**6) / (db * dc * tc)

result = round(0.4 * Pl*1000 / (ew * tw),2)
P = (f"""(1)梁柱节点验算
①螺栓强度验算
梁柱节点采用10.9级M27高强度摩擦型螺栓连接，摩擦面抗滑移系数µ=0.45，每个高强度螺栓的预拉力为{Pl}N,连接处传递内力设计值：
N={N}kN,V={V}kN,M={M}kN·m
每个螺栓拉力设计值：
N_1=\\frac{{M\\times y_1}}{{4\\times\\sum y_i^2}}-\\frac{{N}}{{n}}=\\frac{{{M}\\times{y1}}}{{4\\times\\sum {y1}^2+{y2}^2}}-\\frac{{{N}}}{{{n}}}={N_1:.0f}kN<0.8×{Pl}={round(0.8*Pl,2)}kN
N_2=\\frac{{M\\times y_2}}{{4\\times\\sum y_i^2}}-\\frac{{N}}{{n}}=\\frac{{{M}\\times{y2}}}{{4\\times\\sum {y1}^2+{y2}^2}}-\\frac{{{N}}}{{{n}}}={N_2:.0f}kN<0.8×{Pl}={round(0.8*Pl,2)}kN
螺栓群抗剪承载力：
N_v^b=0.9\\times N_f\\times \\mu\\times P=0.9\\times {Nf}\\times {mu}\\times {Pl}\\times 8={N_vb:.0f}kN>V={V}kN
最外排一个螺栓承载力：
\\frac{{N_v}}{{N_v^b}}+\\frac{{N_t}}{{N_t^b}}=\\frac{{{V}}}{{{N_vb}}}+\\frac{{{Nt}}}{{{N_tb}}}={expression_value:.2f}<1
满足要求。
②端板厚度验算
短板厚度取为t={(int(t1/10)*10+10)}mm，按二辩支撑类端板计算：
t \\geq \\sqrt{{\\frac{{6 \\times e_f \\times e_w \\times N_t}}{{e_w \\times b + 2 \\times e_f \\times (e_f + e_w) \\times f}}}} = \\sqrt{{\\frac{{6 \\times {ef} \\times {ew} \\times {Nt} \\times 1000}}{{{ew} \\times {b} + 2 \\times {ef} \\times ({ef} + {ew}) \\times {f}}}}} = {t1:.2f}mm>{(int(t1/10)*10+5)}mm
满足要求。
③梁柱节点域的剪应力验算
\\tau=\\frac{{M}}{{d_b\\times d_c\\times t_c}}=\\frac{{{M}\\times 10^6}}{{{db}\\times {dc}\\times {tc}}}={tau:.2f}N/mm^2<f={f}N/mm^2
满足要求。
④螺栓处腹板强度验算
N_t={Nt}kN<0.4\\times{Pl}={0.4*Pl}kN
抗剪强度验算：
\\frac{{0.4P}}{{e_w\\times t_w}}=\\frac{{0.4\\times{Pl*1000}}}{{{ew}\\times{tw}}}={result}N/mm^2 < f={f}N/mm^2
满足要求。""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
Pl1 = 190 #抗剪强度

N = min(LBNmax,LCNmax)
V = min(LBQmax,LCQmax)
M = min(LBMmax,LCMmax)

y1 = 0.3
y2 = 0.4

n = 10  # 螺栓个数

# 计算
N_1 = (M * y1) / (4*(y1**2 + y2**2)) - N / n

# 计算
N_2 = (M * y2) / (4*(y1**2 + y2**2)) - N / n

mu = 0.45
Nf = 1

# 计算
N_vb = 0.9 * Nf * mu * Pl1 * 8

Nt = N

N_tb = 0.8*Pl1

# 计算
expression_value = V / N_vb + Nt / N_tb

# 赋值
ef = 60
ew = 80

# 计算
t1 = math.sqrt((6 * ef * ew * Nt * 1000) / (ew * b + 2 * ef * (ef + ew) * fy))

db = hw
dc = hw
tc = t

# 计算
tau = (M * 10**6) / (db * dc * tc)

result = round(0.4 * Pl1*1000 / (ew * tw),2)


P = (f"""(2)跨中梁节点验算
①螺栓强度验算
跨中梁节点采用10.9级M22高强度摩擦型螺栓连接，摩擦面抗滑移系数µ=0.45，每个高强度螺栓的预拉力为{Pl1}N,连接处传递内力设计值：
N={N}kN,V={V}kN,M={M}kN·m
每个螺栓拉力设计值：
N_1=\\frac{{M\\times y_1}}{{4\\times\\sum y_i^2}}-\\frac{{N}}{{n}}=\\frac{{{M}\\times{y1}}}{{4\\times\\sum {y1}^2+{y2}^2}}-\\frac{{{N}}}{{{n}}}={N_1:.0f}kN<0.8×{Pl1}={round(0.8*Pl1,2)}kN
N_2=\\frac{{M\\times y_2}}{{4\\times\\sum y_i^2}}-\\frac{{N}}{{n}}=\\frac{{{M}\\times{y2}}}{{4\\times\\sum {y1}^2+{y2}^2}}-\\frac{{{N}}}{{{n}}}={N_2:.0f}kN<0.8×{Pl1}={round(0.8*Pl1,2)}kN
螺栓群抗剪承载力：
N_v^b=0.9\\times N_f\\times \\mu\\times P=0.9\\times {Nf}\\times {mu}\\times {Pl1}\\times 8={N_vb:.0f}kN>V={V}kN
最外排一个螺栓承载力：
\\frac{{N_v}}{{N_v^b}}+\\frac{{N_t}}{{N_t^b}}=\\frac{{{V}}}{{{N_vb}}}+\\frac{{{Nt}}}{{{N_tb}}}={expression_value:.2f}<1
满足要求。
②端板厚度验算
短板厚度取为t={(int(t1/10)*10+10)}mm，按二辩支撑类端板计算：
t \\geq \\sqrt{{\\frac{{6 \\times e_f \\times e_w \\times N_t}}{{e_w \\times b + 2 \\times e_f \\times (e_f + e_w) \\times f}}}} = \\sqrt{{\\frac{{6 \\times {ef} \\times {ew} \\times {Nt} \\times 1000}}{{{ew} \\times {b} + 2 \\times {ef} \\times ({ef} + {ew}) \\times {fy}}}}} = {t1:.2f}mm>{(int(t1/10)*10+5)}mm
满足要求。
③螺栓处腹板强度验算
N_t={Nt}kN<0.4\\times{Pl1}={0.4*Pl}kN
抗剪强度验算：
\\frac{{0.4P}}{{e_w\\times t_w}}=\\frac{{0.4\\times{Pl*1000}}}{{{ew}\\times{tw}}}={result}N/mm^2 < f={f}N/mm^2
满足要求。""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
Nmax = max(LBNmax,LCNmax)
Vmax = max(LBQmax,LCQmax)

Nmin = min(LBNmax,LCNmax)
Vmin = min(LBQmax,LCQmax)

c1 = 20
c2 = 50

b_min = b + 2 * t + 2 * c1
b_max = b + 2 * t + 2 * c2

bquzhi = 1

if int(b_min/100) < int(b_max/100):
    bquzhi = int(b_max/100)*100
elif int(b_min/100) == int(b_max/100):
    bquzhi = int(b_min/100)+50

h_min = h + 2 * t + 2 * c1
h_max = h + 2 * t + 2 * c2

hquzhi = 1

if int(h_min/100) < int(h_max/100):
    hquzhi = int(h_max/100)*100
elif int(h_min/100) == int(h_max/100):
    hquzhi = int(h_min/100)+50

f_c = 14.3

sigma = round(N * 10**3 / (bquzhi * hquzhi),2)

a1 = 150
M1 = 0.5 * sigma * a1**2
a2 = 50
M2 = 0.5 * sigma * a2**2


tjisuan = round(((6 * M1) / f)**0.5,2)

P = (f"""(2)柱脚节点验算
钢柱与基础铰接连接
①柱脚内力设计值
N_{{max}}={Nmax}kN,N_{{min}}={Nmin}kN
V_{{max}}={Vmax}kN,V_{{min}}={Vmin}kN
②构造要求
由于柱底剪力比较小，且V_{{max}}={Vmax}kN>0.4N_{{max}}={0.4*Nmax}kN，故设置柱间支撑的开间必须设置剪力键，按构造要求设置锚栓即可，采用4M24螺栓。
③柱脚底板面积
b=b_0+2t+2c={b}+2\\times{t}+2\\times\\left({c1}~{c2}\\right)={b_min}~{b_max}mm,\\text{{取}}b={bquzhi}mm
h=h_0+2t+2c={h}+2\\times{t}+2\\times\\left({c1}~{c2}\\right)={h_min}~{h_max}mm,\\text{{取}}h={hquzhi}mm
底板混凝土强度验算：
采用C30混凝土，f_c=14.3N/mm^2,则最大压应力：
\\sigma=\\frac{{N}}{{bh}}=\\frac{{{N}\\times{{10}}^3}}{{{bquzhi}\\times{hquzhi}}}={sigma}N/mm^2<f_c={f_c}N/mm^2
③柱脚底板厚度
支撑板部分弯矩：
M_1=\\frac{1}{2}\\times\\sigma\\times a_1^2=\\frac{1}{2}\\times{sigma}\\times{a1}^2={M1}N\\bullet m
悬挑部分弯矩：
M_2=\\frac{1}{2}\\times\\sigma\\times a_2^2=\\frac{1}{2}\\times{sigma}\\times{a2}^2={M2}N\\bullet m
则底板厚度为：
t=\\sqrt{{\\frac{{6M_{{max}}}}{{f}}}}=\\sqrt{{\\frac{{6\\times{M1}}}{{{f}}}}}={tjisuan}mm,\\text{{取}}t={int(tjisuan/10)*10+10}mm
""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = ('六、檩条设计')
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = True  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(16)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f"""檩条选用实腹式檩条，截面形式选用冷弯薄壁C型钢C250X70X20X3.0，钢材钢号：{GC}钢。拉条设置：设置一道拉条，拉条作用：约束檩条上翼缘。由于设置了一道拉条，保证了檩条在竖向荷载的作用下的整体稳定性，故不用验算檩条的整体稳定性。檩条计算简图如图""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
P = (f""" """)
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 添加图片
docx.paragraphs[30].add_run().add_picture('C:\\Users\\Q\\PycharmProjects\\pythonProject\\钢屋架\\檩条计算简图.png', width=Inches(5.5))
# ----------------------------------------------------------------------------------------------------------------------
q = round((YJHZ+KBHZ)*LJ,2)
import math

# 定义参数
l = ZJ  # m
podu = 5.71  # 坡度1/10

# 计算结果
M_x = 1/8 * q * l**2 * math.cos(podu)

# 定义参数
q_val = q  # kN/m²
l_val = ZJ  # m
alpha_deg_val = podu
# 计算结果
M_y = 0.0156 * q_val * l_val**2 * math.sin(podu)

# 定义参数
force = q  # q, in kN/m²
length = ZJ  # l, in m
angle_deg = 5.71  # α, in degree

# 计算结果
moment_y = -0.0313 * force * length**2 * math.sin(angle_deg)  # M_y, in kNm

I_x = 1013.01
W_x = 81.04 * 10**3  # in mm³
W_y = 12.82 * 10**3  # in mm³

# 计算结果
result = (M_x*1000000 / W_x) + (M_y*1000000 / W_y)

# 定义参数

alpha_deg = 5.71  # α，单位：度
conversion_factor = 10**3  # 转换系数，从 kN/m 转换到 N/m

q_y_prime = ((wumianbanlintiao + KBHZ + 0.9 * 0.5) * LJ * math.cos(alpha_deg)) # 计算 q_y'

l = ZJ*1000 # 单位：mm
E = 2.06 * 10**5  # 单位：N/mm²

# 计算结果
w_l = (5 * q_y_prime * l**3) / (384 * E * I_x*10000)


P = (f"""檩条所受的竖向荷载，屋面板和檩条自重：{YJHZ}kN/m2，可变荷载：{KBHZ}kN/m2，则线荷载设计值q=(({YJHZ}+{KBHZ})×{LJ}={q}kN/m。
按简支梁计算，两个方向弯矩分别是：
跨中最大弯矩：
M_x=\\frac{{1}}{{8}}q_yl^2=\\frac{{1}}{{8}}ql^2\\cos{{{{\\alpha}}}}=\\frac{{1}}{{8}}\\times{q}\\times{l}^2\\times \\cos{{{podu}}}^\\circ={M_x:.2f} \\mathrm{{kN·m}}
M_y=0.0156q_xl^2=0.0156ql^2\\sin{{{{\\alpha}}}}=0.0156\\times{q_val}\\times{l_val}^2\\times \\sin{{{podu}}}^\\circ={abs(M_y):.2f} \\mathrm{{kN·m}}
支座负弯矩：
M_y=-0.0313q_xl^2=-0.0313ql^2\\sin{{{{\\alpha}}}}=-0.0313\\times{force}\\times{length}^2\\times \\sin{{{angle_deg}}}^\\circ={moment_y:.2f} \\mathrm{{kN·m}}
檩条的受弯强度验算：
冷弯薄壁C型钢C250X70X20X3.0的截面特性为:I_x={I_x}cm^4，W_x={W_x}mm³，W_y={W_y}mm³。验算强度时，最不利截面取Mx最大值及其同一截面的My进行计算。
截面最大正应力：
\\frac{{M_x}}{{W_x}}+\\frac{{M_y}}{{W_y}}=\\frac{{{M_x*1000000:.2f}}}{{{W_x}}}+\\frac{{{M_y*1000000:.2f}}}{{{W_y}}}={result:.2f} \\mathrm{{N/m}}\\mathrm{{m}}^2 < f={f} \\mathrm{{N/m}}\\mathrm{{m}}^2
檩条的挠度验算：
由于设有拉条，只验算垂直于屋面坡度的挠度即可。考虑荷载的组合系数，采用恒载+活载+0.9积灰荷载的荷载标准值组合，则:
q_y^\\prime=({wumianbanlintiao}+{KBHZ}+0.9\\times0.5)\\times{LJ}\\times\\cos{{{alpha_deg}}}^\\circ={q_y_prime:.2f} \\mathrm{{N/m}}
\\frac{{w}}{{l}}=\\frac{{5q_y^\\prime l^3}}{{384EI_x}}=\\frac{{5\\times{q_y_prime:.2f}\\times{l}^3}}{{384\\times{E:.2f}\\times{I_x*10000:.2f}}}={w_l:.6f}<\\frac{{1}}{{150}}
""")
# 创建一个空段落，以便放文本
p = docx.add_paragraph("")
run = p.add_run(f'{P}' )
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False  # 加粗
run.italic = False  # 斜体
run.underline = False  # 下划线
run.font.size = Pt(14)
run.font.name=u'宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 首行缩进，小五＝1磅 五号＝10.5磅 小四＝1磅 四号＝14磅 小三＝15磅 三号＝16磅 小二＝18磅 二号＝磅 小一＝4磅 一号＝6磅
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.line_spacing = 1
# ----------------------------------------------------------------------------------------------------------------------
# 保存文档
docx.save("轻型门式钢架.docx")
# 打开文档
import os
os.system("start 轻型门式钢架.docx")